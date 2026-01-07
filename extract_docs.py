from docx import Document

files = [
    'Business Concept.docx',
    '90-Day Action Plan - AI HIV Automation (Western Uganda).docx',
    'Client Outreach Plan - Western Uganda (10 First Customers).docx'
]

for file in files:
    print(f'\n{"="*80}')
    print(f'FILE: {file}')
    print(f'{"="*80}\n')
    try:
        doc = Document(file)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        for table in doc.tables:
            for row in table.rows:
                print(' | '.join([cell.text for cell in row.cells]))
    except Exception as e:
        print(f'Error: {e}')
